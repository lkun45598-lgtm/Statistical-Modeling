#!/usr/bin/env python
from __future__ import annotations

import argparse
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


DEFAULT_TEX = Path("paper/main.tex")
DEFAULT_OUT = Path("paper/build/作品全文-组别-作品编号.docx")
TITLE = "基于高分辨率海温资料的南海增暖及海洋热浪风险统计建模研究"


@dataclass
class TableSpec:
    caption: str
    table_path: Path


@dataclass
class FigureSpec:
    caption: str
    image_paths: list[Path]


def _set_run_font(run, font_name: str = "宋体", size: float | None = 12.0, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def _set_paragraph_format(paragraph, first_line: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(24)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)


def _set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for name, size, font, bold in [
        ("Heading 1", 15, "黑体", True),
        ("Heading 2", 14, "楷体", False),
        ("Heading 3", 12, "宋体", True),
    ]:
        style = doc.styles[name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold


def _set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def _add_toc_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "请在 Word 中右键更新域生成目录"
    fld_sep.append(fld_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _clean_latex(text: str) -> str:
    replacements = {
        r"--": "—",
        r"\%": "%",
        r"\,": "",
        r"\cdot": "·",
        r"\times": "×",
        r"\circ": "°",
        r"\alpha": "α",
        r"\beta": "β",
        r"\varepsilon": "ε",
        r"\varphi": "φ",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\bar": "",
        r"\sum": "Σ",
        r"\frac": "",
        r"\textwidth": "",
        r"\noindent": "",
        r"\_": "_",
    }
    out = text
    out = re.sub(r"\\cite\{[^}]+\}", "", out)
    out = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\texttt\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\textbf\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\^\{\\circ\}", "°", out)
    out = out.replace(r"^\circ", "°")
    out = re.sub(r"\^\{([^{}]+)\}", r"^\1", out)
    out = re.sub(r"_\{([^{}]+)\}", r"_\1", out)
    out = out.replace("$", "")
    out = out.replace(r"\(", "").replace(r"\)", "")
    for old, new in replacements.items():
        out = out.replace(old, new)
    out = re.sub(r"\\[a-zA-Z]+", "", out)
    out = out.replace("{", "").replace("}", "")
    out = re.sub(r"\s+", " ", out).strip()
    return out


def _extract_between(tex: str, start: str, end: str) -> str:
    s = tex.index(start) + len(start)
    e = tex.index(end, s)
    return tex[s:e].strip()


def _extract_abstract(tex: str) -> tuple[list[str], str]:
    block = _extract_between(tex, r"\section*{摘要}", r"\newpage")
    keyword_match = re.search(r"\\textbf\{关键词：\}\s*(.+)", block, flags=re.S)
    keywords = _clean_latex(keyword_match.group(1)) if keyword_match else ""
    body = block[: keyword_match.start()] if keyword_match else block
    paragraphs = [_clean_latex(p) for p in re.split(r"\n\s*\n", body) if _clean_latex(p)]
    return paragraphs, keywords


def _extract_body(tex: str) -> str:
    start = tex.index(r"\section{引言}")
    end = tex.index(r"\ifanonymous\else", start)
    return tex[start:end].strip()


def _parse_table_file(path: Path) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith(r"\begin") or line.startswith(r"\end") or line.startswith(r"\hline"):
            continue
        line = line.rstrip("\\").strip()
        if not line:
            continue
        rows.append([_clean_latex(cell.strip()) for cell in line.split("&")])
    return rows


def _add_word_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(width):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_run_font(run, size=10.5, bold=(i == 0))


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(p, first_line=False)
    run = p.add_run(text)
    _set_run_font(run, size=12)


def _find_table_input(block: str) -> Path | None:
    input_match = re.search(r"\\input\{([^{}]+)\}", block)
    if input_match:
        return Path(input_match.group(1))
    macro_match = re.search(r"\\papertable(?:\[[^\]]+\])?\{([^{}]+)\}", block)
    if macro_match:
        return Path(macro_match.group(1))
    return None


def _add_centered_title(doc: Document, text: str, size: float, font: str = "宋体", bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(p, first_line=False)
    run = p.add_run(text)
    _set_run_font(run, font_name=font, size=size, bold=bold)


def _add_body_paragraph(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    _set_paragraph_format(p, first_line=True)
    run = p.add_run(text)
    _set_run_font(run)


def _add_heading(doc: Document, level: int, title: str) -> None:
    heading = doc.add_heading(level=min(level, 3))
    heading.text = ""
    run = heading.add_run(title)
    if level == 1:
        _set_run_font(run, font_name="黑体", size=15, bold=True)
    elif level == 2:
        _set_run_font(run, font_name="楷体", size=14)
    else:
        _set_run_font(run, font_name="宋体", size=12, bold=True)
    _set_paragraph_format(heading, first_line=False)


def _chinese_number(num: int) -> str:
    numerals = "零一二三四五六七八九十"
    if 1 <= num <= 10:
        return numerals[num]
    if 11 <= num <= 19:
        return "十" + numerals[num - 10]
    tens, ones = divmod(num, 10)
    return numerals[tens] + "十" + (numerals[ones] if ones else "")


def _collect_lists(tex: str) -> tuple[list[TableSpec], list[FigureSpec]]:
    tables: list[TableSpec] = []
    figures: list[FigureSpec] = []
    for match in re.finditer(r"\\begin\{table\}.*?\\end\{table\}", tex, flags=re.S):
        block = match.group(0)
        caption = _clean_latex(re.search(r"\\caption\{([^{}]+)\}", block).group(1))
        table_path = _find_table_input(block)
        if table_path:
            tables.append(TableSpec(caption=caption, table_path=table_path))
    for match in re.finditer(r"\\begin\{figure\}.*?\\end\{figure\}", tex, flags=re.S):
        block = match.group(0)
        captions = re.findall(r"\\caption\{([^{}]+)\}", block)
        caption = _clean_latex(captions[-1]) if captions else "图片"
        image_paths = [Path("paper/figures") / name for name in re.findall(r"\\includegraphics(?:\[[^\]]+\])?\{([^{}]+)\}", block)]
        figures.append(FigureSpec(caption=caption, image_paths=image_paths))
    return tables, figures


def _add_front_matter(doc: Document, tex: str, tables: list[TableSpec], figures: list[FigureSpec]) -> None:
    # Cover page.
    _add_centered_title(doc, "作品编号：TJJM2026XXXXXXX", 16, "黑体", True)
    for _ in range(5):
        doc.add_paragraph()
    _add_centered_title(doc, "2026年（第十二届）全国大学生统计建模大赛", 26, "宋体", False)
    _add_centered_title(doc, "参 赛 作 品", 26, "宋体", False)
    for _ in range(4):
        doc.add_paragraph()
    _add_centered_title(doc, "参赛学校：华南农业大学", 22, "仿宋", False)
    _add_centered_title(doc, "论文题目：", 22, "仿宋", False)
    _add_centered_title(doc, TITLE, 16, "宋体", False)
    _add_centered_title(doc, "参赛队员：待填写", 22, "仿宋", False)
    _add_centered_title(doc, "指导老师：（暂时不填）", 22, "仿宋", False)
    doc.add_page_break()

    _add_centered_title(doc, TITLE, 16, "宋体", False)
    _add_heading(doc, 1, "摘要")
    abstract, keywords = _extract_abstract(tex)
    for para in abstract:
        _add_body_paragraph(doc, para)
    p = doc.add_paragraph()
    _set_paragraph_format(p, first_line=False)
    run = p.add_run(f"关键词：{keywords}")
    _set_run_font(run, font_name="黑体", size=12, bold=True)
    doc.add_page_break()

    _add_heading(doc, 1, "目录")
    p = doc.add_paragraph()
    _add_toc_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()

    _add_heading(doc, 1, "表格与插图清单")
    _add_body_paragraph(doc, "表格清单：")
    for i, spec in enumerate(tables, start=1):
        _add_body_paragraph(doc, f"表{i}  {spec.caption}")
    _add_body_paragraph(doc, "插图清单：")
    for i, spec in enumerate(figures, start=1):
        _add_body_paragraph(doc, f"图{i}  {spec.caption}")
    doc.add_page_break()


def _process_table_block(doc: Document, block: str, table_no: int) -> None:
    caption_match = re.search(r"\\caption\{([^{}]+)\}", block)
    table_path = _find_table_input(block)
    if not caption_match or not table_path:
        return
    caption = _clean_latex(caption_match.group(1))
    rows = _parse_table_file(table_path)
    _add_caption(doc, f"表{table_no}  {caption}")
    _add_word_table(doc, rows)


def _process_figure_block(doc: Document, block: str, figure_no: int) -> None:
    captions = re.findall(r"\\caption\{([^{}]+)\}", block)
    caption = _clean_latex(captions[-1]) if captions else "图片"
    image_names = re.findall(r"\\includegraphics(?:\[[^\]]+\])?\{([^{}]+)\}", block)
    paths = [Path("paper/figures") / name for name in image_names]
    for path in paths:
        if not path.exists():
            continue
        width = Inches(5.8) if len(paths) == 1 else Inches(4.8)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=width)
    _add_caption(doc, f"图{figure_no}  {caption}")


def _add_references(doc: Document, body: str) -> None:
    refs_match = re.search(r"\\begin\{thebibliography\}\{99\}(.*?)\\end\{thebibliography\}", body, flags=re.S)
    if not refs_match:
        return
    _add_heading(doc, 1, "参考文献")
    refs = re.findall(r"\\bibitem\{[^}]+\}\s*(.*?)(?=\\bibitem\{|$)", refs_match.group(1), flags=re.S)
    for idx, raw in enumerate(refs, start=1):
        text = _clean_latex(raw.replace("\n", " "))
        _add_body_paragraph(doc, f"[{idx}] {text}")


def _add_body_from_latex(doc: Document, tex: str) -> None:
    body = _extract_body(tex)
    body_without_refs = re.split(r"\\section\*\{参考文献\}", body, maxsplit=1)[0]

    table_no = 0
    figure_no = 0
    section_no = 0
    subsection_no = 0
    pos = 0
    token_re = re.compile(
        r"\\section\{([^{}]+)\}|\\subsection\{([^{}]+)\}|\\section\*\{([^{}]+)\}|"
        r"\\begin\{table\}.*?\\end\{table\}|\\begin\{figure\}.*?\\end\{figure\}|"
        r"\\begin\{equation\}.*?\\end\{equation\}|\\begin\{itemize\}.*?\\end\{itemize\}",
        flags=re.S,
    )

    def flush_text(chunk: str) -> None:
        cleaned = chunk.strip()
        if not cleaned:
            return
        paragraphs = re.split(r"\n\s*\n", cleaned)
        for para in paragraphs:
            text = _clean_latex(para)
            if text:
                _add_body_paragraph(doc, text)

    for match in token_re.finditer(body_without_refs):
        flush_text(body_without_refs[pos : match.start()])
        block = match.group(0)
        if match.group(1):
            section_no += 1
            subsection_no = 0
            _add_heading(doc, 1, f"{_chinese_number(section_no)}、{_clean_latex(match.group(1))}")
        elif match.group(2):
            subsection_no += 1
            _add_heading(doc, 2, f"（{_chinese_number(subsection_no)}）{_clean_latex(match.group(2))}")
        elif match.group(3):
            _add_heading(doc, 1, _clean_latex(match.group(3)))
        elif block.startswith(r"\begin{table}"):
            table_no += 1
            _process_table_block(doc, block, table_no)
        elif block.startswith(r"\begin{figure}"):
            figure_no += 1
            _process_figure_block(doc, block, figure_no)
        elif block.startswith(r"\begin{equation}"):
            equation = re.sub(r"\\begin\{equation\}|\\end\{equation\}", "", block).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_format(p, first_line=False)
            run = p.add_run(_clean_latex(equation))
            _set_run_font(run)
        elif block.startswith(r"\begin{itemize}"):
            items = re.findall(r"\\item\s*(.*?)(?=\\item|\\end\{itemize\})", block, flags=re.S)
            for item in items:
                p = doc.add_paragraph(style=None)
                _set_paragraph_format(p, first_line=False)
                run = p.add_run("• " + _clean_latex(item))
                _set_run_font(run)
        pos = match.end()

    flush_text(body_without_refs[pos:])
    _add_references(doc, body)


def _add_appendix_and_ack(doc: Document) -> None:
    _add_heading(doc, 1, "附录：可复现计算流程")
    _add_body_paragraph(doc, "本文主要计算脚本均保存在项目仓库的 scripts/ 目录下，核心流程如下：")
    for item in [
        "南海 OSTIA 裁剪：20_prepare_ostia_scs.py；",
        "外部气候指数和风场下载：21_download_external_drivers.py；",
        "月尺度 SSTA 和趋势产品：30_build_monthly_sst_products.py；",
        "月尺度图件：31_make_monthly_sst_figures.py；",
        "日尺度海洋热浪指标：32_build_daily_mhw_products.py；",
        "MHW 图件和趋势表：33_make_mhw_figures.py；",
        "驱动因子解释分析：34_build_driver_analysis.py；",
        "Word 论文生成：35_build_word_paper.py；",
        "分区统计：36_build_subregion_analysis.py；",
        "稳健性和阈值敏感性：37_build_robustness_analysis.py；",
        "综合海洋热风险指数：38_build_heat_risk_index.py；",
        "最终论文图表资产清单：39_make_paper_tables_figures.py。",
    ]:
        p = doc.add_paragraph()
        _set_paragraph_format(p, first_line=False)
        run = p.add_run("• " + item)
        _set_run_font(run)
    _add_body_paragraph(
        doc,
        "主要本地数据包括原始/裁剪后日尺度数据 ostia_scs_daily.zarr、月尺度数据 ostia_scs_monthly.zarr，以及 scs_5s25n/analysis/ 下的统计分析输出。日尺度 MHW 计算采用纬度块并行方式完成，主模型阈值为 90% 分位，稳健性检验中额外计算 85% 和 95% 分位阈值。所有缺失月份保持缺失状态，不进行插值。",
    )
    _add_heading(doc, 1, "致谢")
    _add_body_paragraph(
        doc,
        "感谢本次竞赛组织方提供统计建模实践平台。感谢指导教师和同学在选题讨论、数据处理和论文写作过程中给予的帮助。本文使用的 OSTIA、NOAA/PSL 气候指数和 NCEP/NCAR 再分析资料均为公开数据，谨向相关数据生产和维护团队表示感谢。",
    )


def build_docx(args: argparse.Namespace) -> None:
    tex_path = Path(args.tex)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tex = tex_path.read_text(encoding="utf-8")

    tables, figures = _collect_lists(tex)
    doc = Document()
    _set_page(doc)
    _set_document_styles(doc)

    _add_front_matter(doc, tex, tables, figures)
    _add_body_from_latex(doc, tex)
    _add_appendix_and_ack(doc)

    doc.save(output_path)
    if args.alias:
        alias = Path(args.alias)
        alias.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(output_path, alias)
    print(f"[write] {output_path}")
    if args.alias:
        print(f"[write] {args.alias}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the full Word version of the competition paper.")
    parser.add_argument("--tex", default=str(DEFAULT_TEX))
    parser.add_argument("--output", default=str(DEFAULT_OUT))
    parser.add_argument("--alias", default="paper/build/full_paper.docx")
    args = parser.parse_args()
    build_docx(args)


if __name__ == "__main__":
    main()
